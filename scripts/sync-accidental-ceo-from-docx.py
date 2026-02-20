#!/usr/bin/env python3
"""
Extract Accidental CEO content from the published book .docx and update
ao-knowledge-hq-kit/knowledge/accidental-ceo/*.md, preserving each file's front matter.
"""
import re
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Run: pip3 install python-docx")
    raise

DOCX = Path("/Users/mediaphish/Downloads/Accidental_CEO_by_Bart-Paden-6-x-9-formatted file.docx")
CORPUS_DIR = Path(__file__).resolve().parent.parent / "ao-knowledge-hq-kit" / "knowledge" / "accidental-ceo"

# (start_para_inclusive, end_para_exclusive) for each section (0-based paragraph index)
def find_sections(doc):
    paras = doc.paragraphs
    sections = []
    # Preface: from "A Note to the Reader" through paragraph before "Chapter 1"
    i = 0
    while i < len(paras):
        if paras[i].text.strip() == "A Note to the Reader":
            preface_start = i
            break
        i += 1
    else:
        raise SystemExit("Could not find 'A Note to the Reader'")
    while i < len(paras):
        p = paras[i]
        if getattr(p.style, 'name', None) == "Heading 1" and p.text.strip() == "Chapter 1":
            sections.append(("preface", preface_start, i))
            break
        i += 1
    # Chapters: each has Heading 1 "Chapter N", then Heading 2 (title), then body until next Heading 1
    idx = i
    while idx < len(paras):
        p = paras[idx]
        if getattr(p.style, 'name', None) != "Heading 1" or not p.text.strip().startswith("Chapter "):
            idx += 1
            continue
        num = re.match(r"Chapter (\d+)", p.text.strip())
        num = int(num.group(1)) if num else 0
        # Next para should be Heading 2 (chapter title)
        title_idx = idx + 1
        if title_idx >= len(paras):
            break
        body_start = title_idx + 1
        # Find next Heading 1 = end of this chapter body
        body_end = body_start
        while body_end < len(paras):
            if getattr(paras[body_end].style, 'name', None) == "Heading 1":
                break
            body_end += 1
        sections.append((f"chapter-{num}", body_start, body_end))
        idx = body_end
    return sections

def extract_body(doc, start, end):
    paras = doc.paragraphs
    parts = []
    for i in range(start, min(end, len(paras))):
        t = paras[i].text.strip()
        if t and t != ".":
            parts.append(t)
    return "\n\n".join(parts)

def get_front_matter(content):
    """Return (front_matter_string, rest_after_front_matter)."""
    if not content.startswith("---"):
        return "", content
    idx = content.index("---", 3)
    return content[: idx + 3], content[idx + 3:].lstrip("\n")

def main():
    doc = Document(str(DOCX))
    sections = find_sections(doc)
    # Map section label to corpus filename
    chapter_files = [
        "chapter-1-the-accidental-part.md",
        "chapter-2-what-he-gave-me.md",
        "chapter-3-the-long-road-alone.md",
        "chapter-4-from-the-beginning.md",
        "chapter-5-building-culture.md",
        "chapter-6-the-cost-of-the-climb.md",
        "chapter-7-the-bridge.md",
        "chapter-8-the-fire-after-the-fall.md",
        "chapter-9-the-bridge.md",
        "chapter-10-the-rise-of-archetype.md",
        "chapter-11-introduction-to-the-fundamentals.md",
        "chapter-12-do-the-work.md",
        "chapter-13-protect-the-culture.md",
        "chapter-14-clarity-beats-chaos.md",
        "chapter-15-empowerment-over-control.md",
        "chapter-16-sacrifice-builds-influence.md",
        "chapter-17-invest-in-people-growth-is-a-shared-climb.md",
        "chapter-18-trust-is-the-currency.md",
        "chapter-19-feedback-fuels-growth.md",
        "chapter-20-feedback-ecosystem.md",
        "chapter-21-shield-and-mirror.md",
        "chapter-22-operationalizing-voice.md",
        "chapter-23-sustainable-servant.md",
        "chapter-24-landing-the-plane.md",
    ]
    label_to_file = {"preface": "a-note-to-the-reader.md"}
    for i, f in enumerate(chapter_files, 1):
        label_to_file[f"chapter-{i}"] = f

    for label, start, end in sections:
        fname = label_to_file.get(label)
        if not fname:
            continue
        path = CORPUS_DIR / fname
        if not path.exists():
            print("Skip (no file):", fname)
            continue
        body = extract_body(doc, start, end)
        content = path.read_text(encoding="utf-8")
        fm, _ = get_front_matter(content)
        new_content = fm + "\n" + body.strip() + "\n"
        path.write_text(new_content, encoding="utf-8")
        print("Updated:", fname)

if __name__ == "__main__":
    main()
